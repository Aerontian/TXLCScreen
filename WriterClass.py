# -*- coding: UTF-8 -*
import os
from docx.enum.text import *
from docx.enum.section import *
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.table import *
import copy
from docx import Document
import random


class Writer:

    def __init__(self):

        # self.filename = "/home/aeront/PycharmProjects/TXLCScreen/FactorTest/waa/Doc/test.docx"

        # self.path_template =  \
        #     "/home/aeront/PycharmProjects/TXLCScreen/FactorTest/waa/Doc/template.docx"
        # self.test_pic_path_filename = \
        #     "/home/aeront/PycharmProjects/TXLCScreen/FactorTest/waa/Doc/test_pic.docx"
        # if os.path.exists(self.path_template):
        #     self.temp_doc = Document(self.path_template)
        # else:
        #     print("The template docx does not exist.")
        #     os._exit(1)

        self.doc = Document()
        # self.doc = copy.deepcopy(self.temp_doc)

        # #################################### A4 Paper #####################################
        current_section = self.doc.sections[0]
        current_section.orientation = WD_ORIENTATION.PORTRAIT
        current_section.page_width = Inches(8.3)
        current_section.page_height = Inches(11.7)
        # current_section.footer.text = u"宁波三角洲资产管理有限公司"
        # #################################### A4 Paper #####################################

    def test_report(self, factor_info, df_ic, df_reg, df_gp):

        factor_id = factor_info[0]
        path_factor = "".join(["./FactorTest", "/", factor_id])
        path_factor_pic = "".join([path_factor, '/', 'Pic'])
        # path_factor_xls = "".join([path_factor, '/', 'Xls'])
        path_factor_doc = "".join([path_factor, '/', 'Doc'])

        index_pic = 1
        index_tab = 1

        factor_id_cn = factor_info[1]
        heading_report = "因子: {0}收益率测试".format(factor_id_cn)
        self.set_heading(heading_report, 1, align_hrz='center')

        heading_ic = "1, ic 测试"
        self.set_heading(heading_ic, 2)

        pic_filename = "".join([path_factor_pic, "/", factor_id, "_", "bar_ic", ".png"])
        self.set_picture(pic_filename)
        cap_txt = "图{0} ic柱状图".format(index_pic)
        self.set_pic_caption(cap_txt)
        index_pic += 1

        cap_txt = "表{0} ic值t检验".format(factor_id_cn)
        self.set_tab_caption(cap_txt)
        self.set_ic_test_tab(df_ic)
        index_tab += 1

        heading_ic = "2, 截面回归测试"
        self.set_heading(heading_ic, 2)

        for ii_index, ii_py in enumerate(df_reg.index.tolist()):

            pic_filename = "".join([path_factor_pic, "/", factor_id, "_", "regression", "_", ii_py, ".png"])
            self.set_picture(pic_filename)
            cap_txt = "图{0} 收益回归测试".format(index_pic)
            self.set_pic_caption(cap_txt)
            index_pic += 1

            cap_txt = "表{0} {1} 收益回归测试指标".format(index_tab, ii_py)
            self.set_tab_caption(cap_txt)
            df_reg_item = df_reg.iloc[ii_index:ii_index + 1, :]
            self.set_regression_test_tab(df_reg_item)
            index_tab += 1

        heading_gp = "3, 分组测试"
        self.set_heading(heading_gp, 2)

        pic_filename = "".join([path_factor_pic, "/", factor_id, "_", "group", ".png"])
        self.set_picture(pic_filename)
        cap_txt = "图{0} 分组测试净值曲线".format(index_pic)
        self.set_pic_caption(cap_txt)
        index_pic += 1

        # self.doc.add_page_break()
        new_section = self.doc.add_section(WD_SECTION.ODD_PAGE)
        new_width, new_height = new_section.page_height, new_section.page_width
        new_section.orientation = WD_ORIENT.LANDSCAPE
        new_section.page_width = new_width
        new_section.page_height = new_height

        cap_txt = "表{0} 分组测试指标".format(index_tab, ii_py)
        self.set_tab_caption(cap_txt)
        self.set_group_test_tab(df_gp)
        index_tab += 1

        doc_filename = "".join([path_factor_doc, "/", factor_id, "_", "test", ".docx"])
        self.save_file(doc_filename)

    def set_heading(self, head_txt, head_level, align_hrz='left'):
        run = self.doc.add_heading("", level=head_level).add_run(head_txt)
        run.font.color.rgb = RGBColor(0, 0, 0)
        pr_formant = self.doc.paragraphs[-1].paragraph_format

        if align_hrz == 'left':
            pr_formant.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif align_hrz == 'right':
            pr_formant.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        else:
            pr_formant.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def set_picture(self, pic_path_filename):

        pic = self.doc.add_picture(pic_path_filename)
        pic.height = int(pic.height * 0.8)
        pic.width = int(pic.width * 0.8)
        last_paragraph = self.doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def set_pic_caption(self, caption_txt):

        run = self.doc.add_paragraph("").add_run(caption_txt)
        run.font.color.rgb = RGBColor(0, 0, 0)
        run.font.bold = True
        paragraph = self.doc.paragraphs[-1]
        paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def set_tab_caption(self, caption_txt):

        run = self.doc.add_paragraph("").add_run(caption_txt)
        run.font.color.rgb = RGBColor(0, 0, 0)
        run.font.bold = True
        paragraph = self.doc.paragraphs[-1]
        paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def set_ic_test_tab(self, df_ic_test):

        row_num, col_num = df_ic_test.shape
        table = self.doc.add_table(rows=row_num+1, cols=col_num, style="Medium Shading 2 Accent 1")

        table.rows[0].height = Cm(1)
        table.rows[1].height = Cm(1)

        table_location_property_dict = dict()

        heading = table.rows[0].cells
        heading_list = ["t值", "单边p值"]
        for ii_value, ii_cell in zip(heading_list, heading):
            propt = dict()
            if isinstance(ii_value, str):
                propt['value'] = ii_value
            else:
                propt['value'] = "{0:+.2f}".format(ii_value)
            propt['bold'] = True
            propt['size'] = 10
            propt['color'] = RGBColor(0, 0, 0)
            propt['para_alignment'] = WD_ALIGN_PARAGRAPH.CENTER
            propt['cell_vertical_alignment'] = WD_ALIGN_VERTICAL.CENTER
            table_location_property_dict[ii_cell] = propt

        ic_test_ret = df_ic_test.iloc[0, :].tolist()
        content = table.rows[1].cells
        for ii_value, ii_cell in zip(ic_test_ret, content):
            propt = dict()
            if isinstance(ii_value, str):
                propt['value'] = ii_value
            else:
                propt['value'] = "{0:+.2f}".format(ii_value)
            propt['bold'] = True
            propt['size'] = 10
            propt['color'] = RGBColor(0, 0, 0)
            propt['para_alignment'] = WD_ALIGN_PARAGRAPH.CENTER
            propt['cell_vertical_alignment'] = WD_ALIGN_VERTICAL.CENTER
            table_location_property_dict[ii_cell] = propt

        self.dump_data_table(table_location_property_dict, table)

    def set_regression_test_tab(self, df_rg_test):

        row_num, col_num = df_rg_test.shape
        table = self.doc.add_table(rows=row_num + 1, cols=col_num, style="Medium Shading 2 Accent 1")

        table.rows[0].height = Cm(1)
        table.rows[1].height = Cm(1)

        table_location_property_dict = dict()

        heading = table.rows[0].cells
        heading_list = ["f值", "F检验p值", "截距项t值", "截距项t检验p值", "一次项t值", "一次项t检验p值"]
        for ii_value, ii_cell in zip(heading_list, heading):
            propt = dict()
            if isinstance(ii_value, str):
                propt['value'] = ii_value
            else:
                propt['value'] = "{0:+.2f}".format(ii_value)
            propt['bold'] = True
            propt['size'] = 10
            propt['color'] = RGBColor(0, 0, 0)
            propt['para_alignment'] = WD_ALIGN_PARAGRAPH.CENTER
            propt['cell_vertical_alignment'] = WD_ALIGN_VERTICAL.CENTER
            table_location_property_dict[ii_cell] = propt

        rg_test_ret = df_rg_test.iloc[0, :].tolist()
        content = table.rows[1].cells
        for ii_value, ii_cell in zip(rg_test_ret, content):
            propt = dict()
            if isinstance(ii_value, str):
                propt['value'] = ii_value
            else:
                propt['value'] = "{0:+.2f}".format(ii_value)
            propt['bold'] = True
            propt['size'] = 10
            propt['color'] = RGBColor(0, 0, 0)
            propt['para_alignment'] = WD_ALIGN_PARAGRAPH.CENTER
            propt['cell_vertical_alignment'] = WD_ALIGN_VERTICAL.CENTER
            table_location_property_dict[ii_cell] = propt

        self.dump_data_table(table_location_property_dict, table)

    def set_group_test_tab(self, df_gp_test):

        row_num, col_num = df_gp_test.shape
        table = self.doc.add_table(rows=row_num + 1, cols=col_num + 1, style="Medium Shading 2 Accent 1")

        for ii_row in table.rows:
            ii_row.height = Cm(0.8)

        table_location_property_dict = dict()

        heading = table.rows[0].cells
        heading_list = ["分组", "交易日", "净值", "年化收益", "年化波动", "夏普比", "盈利日",
                        "亏损日", "胜率", "盈利平均", "亏损平均", "盈亏比"]
        for ii_value, ii_cell in zip(heading_list, heading):
            propt = dict()
            if isinstance(ii_value, str):
                propt['value'] = ii_value
            else:
                propt['value'] = "{0:+.2f}".format(ii_value)
            propt['bold'] = True
            propt['size'] = 10
            propt['color'] = RGBColor(0, 0, 0)
            propt['para_alignment'] = WD_ALIGN_PARAGRAPH.CENTER
            propt['cell_vertical_alignment'] = WD_ALIGN_VERTICAL.CENTER
            table_location_property_dict[ii_cell] = propt

        heading = table.columns[0].cells[1:]
        heading_list = ["1", "2", "3", "4", "5", "6",
                        "7", "8", "9", "10", "市场"]
        for ii_value, ii_cell in zip(heading_list, heading):
            propt = dict()
            if isinstance(ii_value, str):
                propt['value'] = ii_value
            else:
                propt['value'] = "{0:+.2f}".format(ii_value)
            propt['bold'] = True
            propt['size'] = 10
            if ii_cell != heading[-1]:
                propt['color'] = RGBColor(0, 0, 0)
            else:
                propt['color'] = RGBColor(255, 0, 0)
            propt['para_alignment'] = WD_ALIGN_PARAGRAPH.CENTER
            propt['cell_vertical_alignment'] = WD_ALIGN_VERTICAL.CENTER
            table_location_property_dict[ii_cell] = propt

        for ii_index, ii_row in enumerate(table.rows[1:]):
            group_test_ret = df_gp_test.iloc[ii_index, :].values.tolist()
            content = ii_row.cells[1:]
            index_item = 0
            for ii_value, ii_cell in zip(group_test_ret, content):
                propt = dict()
                if index_item in [0, 5, 6]:
                    propt['value'] = str(int(ii_value))
                else:
                    propt['value'] = "{0:+.4f}".format(ii_value)
                index_item += 1
                # if isinstance(ii_value, str):
                #     propt['value'] = ii_value
                # else:
                #     propt['value'] = "{0:+.2f}".format(ii_value)
                propt['bold'] = False
                propt['size'] = 10
                if ii_index == row_num - 1:
                    propt['color'] = RGBColor(255, 0, 0)
                else:
                    propt['color'] = RGBColor(0, 0, 0)
                propt['para_alignment'] = WD_ALIGN_PARAGRAPH.CENTER
                propt['cell_vertical_alignment'] = WD_ALIGN_VERTICAL.CENTER
                table_location_property_dict[ii_cell] = propt

        self.dump_data_table(table_location_property_dict, table)

    @staticmethod
    def dump_data_table(table_location_property_dict, table):

        # table.columns[0].width = Cm(1)

        for ii_location, ii_propt in table_location_property_dict.items():

            if isinstance(ii_location, tuple):
                cell = table.cell(ii_location[0], ii_location[1])
            else:
                cell = ii_location

            value = ii_propt.get('value', '')
            bold = ii_propt.get('bold')
            para__alignment = ii_propt.get('para_alignment')
            cell_vertical_alignment = ii_propt.get('cell_vertical_alignment')

            # cell.text = value
            run = cell.paragraphs[0].add_run(value)

            cell.paragraphs[0].paragraph_format.alignment = para__alignment
            cell.vertical_alignment = cell_vertical_alignment

            run_font = run.font
            run_font.size = Pt(ii_propt.get('size'))
            run_font.color.rgb = ii_propt.get('color')

            if bold:
                run_font.bold = bold

    def test_pic(self):

        self.doc.add_page_break()
        new_section = self.doc.add_section(WD_SECTION.ODD_PAGE)
        section = self.doc.sections[-1]
        new_width, new_height = section.page_height, section.page_width
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = new_width
        section.page_height = new_height
        # section.orientation = WD_ORIENT.LANDSCAPE
        table = self.doc.add_table(rows=12, cols=12)

        table.style = "Medium Shading 2 Accent 1"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        heading = table.rows[0].cells
        heading_list = ["分组", "交易日",	"净值",	"年化收益", "年化波动", 	"夏普比",
                        "盈利日", "亏损日",	"胜率",	"盈利平均", "亏损平均",	"盈亏比"]

        for ii_row in table.rows:
            ii_row.height = Cm(1)

        for ii_index, ii_item in enumerate(heading_list):
            cell = heading[ii_index]

            value = ii_item
            bold = True
            para__alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell_vertical_alignment = WD_ALIGN_VERTICAL.CENTER

            # cell.text = value
            run = cell.paragraphs[0].add_run(value)

            cell.paragraphs[0].paragraph_format.alignment = para__alignment
            cell.vertical_alignment = cell_vertical_alignment

            font = run.font
            font.size = Pt(10.5)
            run.bold = bold

        heading = table.columns[0].cells[1:]
        heading_list = ["1", "2", "3", "4", "5", "6",
                        "7", "8", "9", "10", "市场"]
        for ii_index, ii_item in enumerate(heading_list):
            cell = heading[ii_index]

            value = ii_item
            bold = True
            para__alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell_vertical_alignment = WD_ALIGN_VERTICAL.CENTER

            # cell.text = value
            run = cell.paragraphs[0].add_run(value)

            cell.paragraphs[0].paragraph_format.alignment = para__alignment
            cell.vertical_alignment = cell_vertical_alignment

            font = run.font
            font.size = Pt(10.5)
            run.bold = bold

        cell_pos_list = [(ii_row, ii_col) for  ii_row in range(1, 12) for ii_col in range(1, 12)]
        for ii_cell_pos in cell_pos_list:

            cell = table.cell(ii_cell_pos[0], ii_cell_pos[1])

            value = "{:+.2f}".format(random.randint(-100,100) * 1.0 /100)
            bold = False
            para__alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell_vertical_alignment = WD_ALIGN_VERTICAL.CENTER

            # cell.text = value
            run = cell.paragraphs[0].add_run(value)

            cell.paragraphs[0].paragraph_format.alignment = para__alignment
            cell.vertical_alignment = cell_vertical_alignment

            font = run.font
            font.size = Pt(10.5)
            run.bold = bold
            if ii_cell_pos[0] == 11:
                font.color.rgb = RGBColor(255, 0, 0)



        # table = self.doc.add_table(rows=2, cols=6)
        #
        # table.style = "Medium Shading 2 Accent 1"
        # table.alignment = WD_TABLE_ALIGNMENT.CENTER
        # # table.autofit = True
        #
        # table_location_property_dict = dict()
        #
        # location = (0, 0)
        # propt = dict()
        # propt['value'] = "f_value"
        # propt['bold'] = True
        # propt['size'] = 10
        # propt['para_alignment'] = WD_ALIGN_PARAGRAPH.CENTER
        # propt['cell_vertical_alignment'] = WD_ALIGN_VERTICAL.CENTER
        # table_location_property_dict[location] = propt
        # table_location_property_dict[location] = propt
        #
        # location = (0, 1)
        # propt = dict()
        # propt['value'] = "F-p_value"
        # propt['bold'] = True
        # propt['size'] = 10
        # propt['para_alignment'] = WD_ALIGN_PARAGRAPH.CENTER
        # propt['cell_vertical_alignment'] = WD_ALIGN_VERTICAL.CENTER
        # table_location_property_dict[location] = propt
        # table_location_property_dict[location] = propt
        #
        # location = (0, 2)
        # propt = dict()
        # propt['value'] = "itcpt:\nt_value"
        # propt['bold'] = True
        # propt['size'] = 10
        # propt['para_alignment'] = WD_ALIGN_PARAGRAPH.CENTER
        # propt['cell_vertical_alignment'] = WD_ALIGN_VERTICAL.CENTER
        # table_location_property_dict[location] = propt
        # table_location_property_dict[location] = propt
        #
        # location = (0, 3)
        # propt = dict()
        # propt['value'] = "itcpt:\nt_p_value"
        # propt['bold'] = True
        # propt['size'] = 10
        # propt['para_alignment'] = WD_ALIGN_PARAGRAPH.CENTER
        # propt['cell_vertical_alignment'] = WD_ALIGN_VERTICAL.CENTER
        # table_location_property_dict[location] = propt
        # table_location_property_dict[location] = propt
        #
        # location = (0, 4)
        # propt = dict()
        # propt['value'] = "linr:\nt_value"
        # propt['bold'] = True
        # propt['size'] = 10
        # propt['para_alignment'] = WD_ALIGN_PARAGRAPH.CENTER
        # propt['cell_vertical_alignment'] = WD_ALIGN_VERTICAL.CENTER
        # table_location_property_dict[location] = propt
        # table_location_property_dict[location] = propt
        #
        # location = (0, 5)
        # propt = dict()
        # propt['value'] = "linr:\nt_p_value"
        # propt['bold'] = True
        # propt['size'] = 10
        # propt['para_alignment'] = WD_ALIGN_PARAGRAPH.CENTER
        # propt['cell_vertical_alignment'] = WD_ALIGN_VERTICAL.CENTER
        # table_location_property_dict[location] = propt
        # table_location_property_dict[location] = propt
        #
        # location = (1, 0)
        # propt = dict()
        # propt['value'] = "6.47"
        # propt['bold'] = True
        # propt['size'] = 10
        # propt['para_alignment'] = WD_ALIGN_PARAGRAPH.CENTER
        # propt['cell_vertical_alignment'] = WD_ALIGN_VERTICAL.CENTER
        # table_location_property_dict[location] = propt
        # table_location_property_dict[location] = propt
        #
        # location = (1, 1)
        # propt = dict()
        # propt['value'] = "0.01"
        # propt['bold'] = True
        # propt['size'] = 10
        # propt['para_alignment'] = WD_ALIGN_PARAGRAPH.CENTER
        # propt['cell_vertical_alignment'] = WD_ALIGN_VERTICAL.CENTER
        # table_location_property_dict[location] = propt
        # table_location_property_dict[location] = propt
        #
        # location = (1, 2)
        # propt = dict()
        # propt['value'] = "1.445"
        # propt['bold'] = True
        # propt['size'] = 10
        # propt['para_alignment'] = WD_ALIGN_PARAGRAPH.CENTER
        # propt['cell_vertical_alignment'] = WD_ALIGN_VERTICAL.CENTER
        # table_location_property_dict[location] = propt
        # table_location_property_dict[location] = propt
        #
        # location = (1, 3)
        # propt = dict()
        # propt['value'] = "0.1485"
        # propt['bold'] = True
        # propt['size'] = 10
        # propt['para_alignment'] = WD_ALIGN_PARAGRAPH.CENTER
        # propt['cell_vertical_alignment'] = WD_ALIGN_VERTICAL.CENTER
        # table_location_property_dict[location] = propt
        # table_location_property_dict[location] = propt
        #
        # location = (1, 4)
        # propt = dict()
        # propt['value'] = "2.5446"
        # propt['bold'] = True
        # propt['size'] = 10
        # propt['para_alignment'] = WD_ALIGN_PARAGRAPH.CENTER
        # propt['cell_vertical_alignment'] = WD_ALIGN_VERTICAL.CENTER
        # table_location_property_dict[location] = propt
        # table_location_property_dict[location] = propt
        #
        # location = (1, 5)
        # propt = dict()
        # propt['value'] = "0.0110"
        # propt['bold'] = True
        # propt['size'] = 10
        # propt['para_alignment'] = WD_ALIGN_PARAGRAPH.CENTER
        # propt['cell_vertical_alignment'] = WD_ALIGN_VERTICAL.CENTER
        # table_location_property_dict[location] = propt
        # table_location_property_dict[location] = propt
        #
        # self.dump_data_table(table_location_property_dict, table)



        # filter_extension = [u"png"]
        # path = "/home/aeront/PycharmProjects/TXLCScreen/FactorTest/waa/Pic"
        # pic_file_list = list()
        #
        # for root, dirs, files in os.walk(path):
        #     for filename in files:
        #         filename_extension = filename.split(".")[1]
        #         if filename_extension in filter_extension:
        #             pic_file_list.append(filename)
        #
        # for ii_pic in pic_file_list:
        #
        #     ii_pic_filename = "".join([path, '/', ii_pic])
        #     pic = self.doc.add_picture(ii_pic_filename)
        #
        #     last_paragraph = self.doc.paragraphs[-1]
        #     last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        #     pic.height = int(pic.height * 0.72)
        #     pic.width = int(pic.width * 0.72)

        # year_list = \
        #     ["2007", "2008", "2009", "2010", "2011", "2012",
        #      "2013", "2014", "2015", "2016", "2017", "2018", "2019"]
        #
        # for ii_index, ii_year in enumerate(year_list):
        #
        #     para = self.doc.add_paragraph()
        #     para_format = para.paragraph_format
        #
        #     headline = "{0}: {1} 收益回归测试\n".format(ii_index, ii_year)
        #     run = para.add_run(headline)
        #     font = run.font
        #
        #     font.name = "SimSun"
        #     font.size = Pt(14)
        #     font.bold =True
        #     para_format.alignment = WD_ALIGN_PARAGRAPH.LEFT




        # document = self.temp_doc
        # table_style = copy.deepcopy(document.tables[-1].style)
        # self.doc.add_table(14, 14, table_style)

    def save_file(self, filename):

        self.doc.save(filename)

    # def dump_data(self):
    #
    #     para = self.doc.paragraphs[0]
    #     run = para.add_run(self.account.trade_day)
    #     font = run.font
    #
    #     font.name = "SimSun"
    #     font.size = Pt(12)
    #     para_format = para.paragraph_format
    #     para_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    #
    #     para = self.doc.add_paragraph()
    #     run = para.add_run(u"\n")
    #     font = run.font
    #
    #     font.name = "SimSun"
    #     font.size = Pt(14)
    #     para_format = para.paragraph_format
    #     para_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    #
    #     para = self.doc.add_paragraph()
    #     para_format = para.paragraph_format
    #     para_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    #     run = para.add_run(self.account.data_source.data_info['client_Id'])
    #     font = run.font
    #     run.bold = True
    #
    #     font.name = "SimSun"
    #     font.size = Pt(14)
    #
    #     table = self.doc.add_table(rows=3, cols=5)
    #     table.style = "Table Grid"
    #     table.alignment = WD_TABLE_ALIGNMENT.CENTER
    #     table.autofit = True
    #     cell = table.cell(1, 0)
    #     other_cell = table.cell(2, 0)
    #     cell.merge(other_cell)
    #
    #     table_location_property_dict = dict()
    #
    #     location = (0, 0)
    #     propt = dict()
    #     propt['value'] = u"产品净值"
    #     propt['bold'] = True
    #     propt['para_alignment'] = WD_ALIGN_PARAGRAPH.CENTER
    #     propt['cell_vertical_alignment'] = WD_ALIGN_VERTICAL.CENTER
    #     table_location_property_dict[location] = propt
    #
    #     location = (0, 1)
    #     propt = dict()
    #     propt['value'] = str(round(self.account.equity_excel[-1, 0], 5)).encode()
    #     propt['para_alignment'] = WD_ALIGN_PARAGRAPH.CENTER
    #     propt['cell_vertical_alignment'] = WD_ALIGN_VERTICAL.CENTER
    #     table_location_property_dict[location] = propt
    #
    #     location = (0, 2)
    #     propt = dict()
    #     propt['value'] = u"净值变化"
    #     propt['bold'] = True
    #     propt['para_alignment'] = WD_ALIGN_PARAGRAPH.CENTER
    #     propt['cell_vertical_alignment'] = WD_ALIGN_VERTICAL.CENTER
    #     table_location_property_dict[location] = propt
    #
    #     location = (0, 3)
    #     propt = dict()
    #     propt['value'] = str(round(self.account.equity_excel[-1, 1], 6)).encode()
    #     propt['bold'] = True
    #     propt['para_alignment'] = WD_ALIGN_PARAGRAPH.CENTER
    #     propt['cell_vertical_alignment'] = WD_ALIGN_VERTICAL.CENTER
    #     table_location_property_dict[location] = propt
    #
    #     location = (1, 0)
    #     propt = dict()
    #     propt['value'] = u"期货"
    #     propt['bold'] = True
    #     propt['para_alignment'] = WD_ALIGN_PARAGRAPH.CENTER
    #     propt['cell_vertical_alignment'] = WD_ALIGN_VERTICAL.CENTER
    #     table_location_property_dict[location] = propt
    #
    #     location = (1, 1)
    #     propt = dict()
    #     propt['value'] = u"保证金"
    #     propt['bold'] = True
    #     propt['para_alignment'] = WD_ALIGN_PARAGRAPH.CENTER
    #     propt['cell_vertical_alignment'] = WD_ALIGN_VERTICAL.CENTER
    #     table_location_property_dict[location] = propt
    #
    #     location = (2, 1)
    #     propt = dict()
    #     propt['value'] = str(self.account.total_margin).encode()
    #     propt['para_alignment'] = WD_ALIGN_PARAGRAPH.CENTER
    #     propt['cell_vertical_alignment'] = WD_ALIGN_VERTICAL.CENTER
    #     table_location_property_dict[location] = propt
    #
    #     location = (1, 2)
    #     propt = dict()
    #     propt['value'] = u"期权净值"
    #     propt['bold'] = True
    #     propt['para_alignment'] = WD_ALIGN_PARAGRAPH.CENTER
    #     propt['cell_vertical_alignment'] = WD_ALIGN_VERTICAL.CENTER
    #     table_location_property_dict[location] = propt
    #
    #     location = (2, 2)
    #     propt = dict()
    #     propt['value'] = str(0).encode()
    #     propt['para_alignment'] = WD_ALIGN_PARAGRAPH.CENTER
    #     propt['cell_vertical_alignment'] = WD_ALIGN_VERTICAL.CENTER
    #     table_location_property_dict[location] = propt
    #
    #     location = (1, 3)
    #     propt = dict()
    #     propt['value'] = u"可用资金"
    #     propt['bold'] = True
    #     propt['para_alignment'] = WD_ALIGN_PARAGRAPH.CENTER
    #     propt['cell_vertical_alignment'] = WD_ALIGN_VERTICAL.CENTER
    #     table_location_property_dict[location] = propt
    #
    #     location = (2, 3)
    #     propt = dict()
    #     propt['value'] = str(round(self.account.cash_available, 2)).encode()
    #     propt['bold'] = True
    #     propt['para_alignment'] = WD_ALIGN_PARAGRAPH.CENTER
    #     propt['cell_vertical_alignment'] = WD_ALIGN_VERTICAL.CENTER
    #     table_location_property_dict[location] = propt
    #
    #     location = (1, 4)
    #     propt = dict()
    #     propt['value'] = u"期货保证金%"
    #     propt['bold'] = True
    #     propt['para_alignment'] = WD_ALIGN_PARAGRAPH.CENTER
    #     propt['cell_vertical_alignment'] = WD_ALIGN_VERTICAL.CENTER
    #     table_location_property_dict[location] = propt
    #
    #     location = (2, 4)
    #     propt = dict()
    #     propt['value'] = str(round(100.0*self.account.total_margin/self.account.end_value, 2)).encode() + u"%"
    #     propt['para_alignment'] = WD_ALIGN_PARAGRAPH.CENTER
    #     propt['cell_vertical_alignment'] = WD_ALIGN_VERTICAL.CENTER
    #     table_location_property_dict[location] = propt
    #
    #     self.dump_data_table(table_location_property_dict, table)
    #
    #     para = self.doc.add_paragraph()
    #     para.add_run(u"\n")
    #
    #     figure_name = self.account.path + u'EquityPng/' + self.account.trade_day + u"_equity.png"
    #     pic = self.doc.add_picture(figure_name)
    #     pic.height = int(pic.height * 0.8)
    #     pic.width = int(pic.width * 0.8)
    #     last_paragraph = self.doc.paragraphs[-1]
    #     last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    #
    #     self.doc.save(self.filename)


if __name__ == "__main__":

    doc_writer = Writer()
    doc_writer.test_pic()
    doc_writer.save_file()




